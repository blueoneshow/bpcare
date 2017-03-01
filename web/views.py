# -*- coding: UTF-8 -*-
from django.shortcuts import render
from django.http import HttpResponse
from django.template import RequestContext
from django.utils import timezone
from django.utils.timezone import localtime
from django.shortcuts import render_to_response, redirect
from django.core.exceptions import ObjectDoesNotExist
from web.models import Message, Diary, Month, Money
from web.forms import MessageForm, DiaryForm, MoneyForm
import StringIO
from docx import *
from docx.shared import Inches
import xlsxwriter

#既有留言總覽
def board(request):
        messages = Message.objects.all()
        response_string = "<a href='/post'>Post</a><hr/>"
        response_string += '<br/>'.join(["user: %s, subject: %s, time: %s" % (q.user, q.subject, q.publication_date) for q in messages])
        return HttpResponse(response_string)

#新增留言
def post(request):
        if request.method == 'POST':
                form = MessageForm(request.POST)
                if form.is_valid():
                        message = Message(user=form.cleaned_data['user'],subject=form.cleaned_data['subject'], publication_date=timezone.now())
                        message.save()
                        return redirect('/board')
        else:
                form = MessageForm()
        return render_to_response('post.html',{'form': form}, context_instance=RequestContext(request))
      
      
# 瀏覽日誌（總覽）
def diary(request):
        diaries = Diary.objects.all().order_by("-id")
        return render_to_response('diary.html', {'diaries': diaries}, context_instance=RequestContext(request))

#瀏覽日誌（分月份）
def diary_month(request, month):
        time_year = int(month)/100
        time_month = int(month)%100
        diaries_m = Diary.objects.filter(time__year=time_year, time__month=time_month).order_by("-id")
        return render_to_response('diary_month.html', {'diaries_m': diaries_m, 'month':month}, context_instance=RequestContext(request))
    
#新增日誌（儲存後轉至所有日誌）
#def diary_add(request):
#        if request.method == 'POST':
#                form = DiaryForm(request.POST)
#                if form.is_valid():
#                        form.save()
#                        year = localtime(timezone.now()).year
#                        month =  localtime(timezone.now()).month
#                        try:
#                                themonth = Month.objects.get(date=year*100+month)
#                        except ObjectDoesNotExist:
#                                themonth = Month(date=year*100+month)
#                                themonth.save()
#                        return redirect("/diary")
#        else:
#                form = DiaryForm()
#        return render_to_response('diary_add.html',{'form': form}, context_instance=RequestContext(request))
 

#新增日誌（儲存後轉至該月份） 
def diary_add(request):
        if request.method == 'POST':
                form = DiaryForm(request.POST)
                if form.is_valid():
                        form.save()
                        year = localtime(timezone.now()).year
                        month =  localtime(timezone.now()).month
                        try:
                                themonth = Month.objects.get(date=year*100+month)
                        except ObjectDoesNotExist:
                                themonth = Month(date=year*100+month)
                                themonth.save()
                        return redirect("/diary/"+str(year*100+month))
        else:
                form = DiaryForm()
        return render_to_response('diary_add.html',{'form': form}, context_instance=RequestContext(request))
      
#首頁
def home(request):
        months = Month.objects.all().order_by("-id")
        return render_to_response('home.html', {'months': months}, context_instance=RequestContext(request))

#該月份日誌匯出至word檔
def diary_word(request, month):
        document = Document()
        docx_title="Diary-"+str(timezone.localtime(timezone.now()).date())+".docx"

        time_year = int(month)/100
        time_month = int(month)%100
        diaries = Diary.objects.filter(time__year=time_year, time__month=time_month).order_by("-id")
        paragraph = document.add_paragraph(u'我的日誌：'+month)
        table = document.add_table(rows=1, cols=2)
        table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = u'日期'
        hdr_cells[1].text = u'內容'
        for diary in diaries:
                row_cells = table.add_row().cells
                row_cells[0].text = str(timezone.localtime(diary.time).strftime("%b %d %Y %H:%M:%S"))
                row_cells[1].text = diary.memo

        # Prepare document for download
        # -----------------------------
        f = StringIO.StringIO()
        document.save(f)
        length = f.tell()
        f.seek(0)
        response = HttpResponse(
                f.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=' + docx_title
        response['Content-Length'] = length

        return response

#月支出
def money(request, month):
        time_year = int(month)/100
        time_month = int(month)%100
        moneys = Money.objects.filter(time__year=time_year, time__month=time_month).order_by("-id")
        return render_to_response('money.html', {'moneys': moneys, 'month':month}, context_instance=RequestContext(request))
      
#新增支出
def money_add(request):
        if request.method == 'POST':
                form = MoneyForm(request.POST)
                if form.is_valid():
                        form.save()
                        year = localtime(timezone.now()).year
                        month =  localtime(timezone.now()).month
                        try:
                                themonth = Month.objects.get(date=year*100+month)
                        except ObjectDoesNotExist:
                                themonth = Month(date=year*100+month)
                                themonth.save()
                        return redirect("/money/"+str(themonth.date))
        else:
                form = MoneyForm()
        return render_to_response('money_add.html',{'form': form}, context_instance=RequestContext(request))
      
      
#匯出
def money_excel(request, month):
        time_year = int(month)/100
        time_month = int(month)%100
        moneys = Money.objects.filter(time__year=time_year, time__month=time_month).order_by("-id")
        output = StringIO.StringIO()
        workbook = xlsxwriter.Workbook(output)
        worksheet = workbook.add_worksheet(month)
        worksheet.write(0,2, u"收縮壓")
        worksheet.write(0,1, u"類別")
        worksheet.write(0,3, u"舒張壓")
        worksheet.write(0,0, u"測量時間")
        counter = 1
        for money in moneys:
                worksheet.write(counter,2, money.item)
                worksheet.write(counter,1, money.kind)
                worksheet.write(counter,3, money.price)
                worksheet.write(counter,0, str(localtime(money.time).strftime("%b %d %Y %H:%M:%S")))
                counter = counter + 1
        workbook.close()
        # xlsx_data contains the Excel file
        response = HttpResponse(content_type='application/vnd.ms-excel')
        response['Content-Disposition'] = 'attachment; filename=Money-'+str(localtime(timezone.now()).date())+'.xlsx'
        xlsx_data = output.getvalue()
        response.write(xlsx_data)
        return response